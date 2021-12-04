import re
from pathlib import Path
import os
import shutil #standard utility
from docx import Document

class FileCombining():
    '''This class combines both Basic and Advanced , Theoretical and Programming Assignments into one
    single word document '''

    def __init__(self, filesList,dest_path):
        self.filesList = filesList
        self.dest_path=dest_path
        self.renamedFiles = []
        self.filterTheoryAssignmentFiles()
        self.filterProgrammingAssignmentFiles()
        self.moveFilesToDestinationFolder()
        self.combineFiles()

    def filterTheoryAssignmentFiles(self):
        '''This function filters only Theoretical-(Basic and Advanced) Assignments and Renames the file
            It appends all Renamed Theory Assignment files to List renamedFiles ( class variable ) '''
        try:
            threoryFilesList = list(filter(lambda x: re.match("[A-Za-z]+_\d.*docx$", x), self.filesList))

            for file in threoryFilesList:
                if file.endswith(').docx'):
                    newFilename1 = re.sub("[A-Za-z]+", "Advanced_Theoretical_Assignment", file.split(' ')[0])
                    os.rename(file, newFilename1 + '.docx')
                    self.renamedFiles.append(newFilename1 + '.docx')
                else:
                    newFilename2 = re.sub("[A-Za-z]+", "Basic_Theoretical_Assignment", file.split('.')[0])
                    os.rename(file, newFilename2 + '.docx')
                    self.renamedFiles.append(newFilename2 + '.docx')

        except Exception as e:
            print("Unable to Filter/Rename Theory Assignment Files",e)

    def filterProgrammingAssignmentFiles(self):
        '''This function filters only Programming-(Basic and Advanced) Assignments and Renames the file
            It appends all Renamed Programming Assignment files to List renamedFiles ( class variable ) '''
        try:
            for file in self.filesList:
                if re.match("[Aa-z]+\d[.]docx$", file):
                    newFilename3 = re.sub("[A-Za-z]+", "Advanced_Programming_Assignment", file.rstrip('docx'))
                    os.rename(file, newFilename3 + 'docx')
                    self.renamedFiles.append(newFilename3 + 'docx')
                elif re.match("[A-Za-z]+_[A-Za-z]+\d[.]docx$", file):
                    newFilename4 = "Basic_" + file
                    os.rename(file, newFilename4)
                    self.renamedFiles.append(newFilename4)

        except Exception as e:
            print("Unable to Filter/Rename Practical Assignment Files:",e)
        self.renamedFiles.sort()

    def moveFilesToDestinationFolder(self):
        '''This Function will move all the renamed files to Destination Folder'''

        path = os.path.join(self.dest_path, "Assignment_Files")
        if not Path(path).is_dir():
            os.mkdir(path)
        for file in self.renamedFiles:
            shutil.move(file, path)
        os.chdir(path)
        self.renamedFiles.sort()

    def combineFiles(self):
        '''This Function combines all the Renamed Assignment files into one single word document'''

        merged_doc=Document()
        try:
            for file in self.renamedFiles:
                doc=Document(file)
                merged_doc.add_heading(file.rstrip('.docx'),1)
                for para in doc.paragraphs:
                    text=para.text
                    merged_doc.add_paragraph(text)
            merged_doc.save('Combined_Assignment_Files.docx')
        except Exception as e:
            print('Unable to combine Files:', e)